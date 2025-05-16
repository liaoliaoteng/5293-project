# summarization_gui.py

import tkinter as tk
from tkinter import ttk, filedialog, scrolledtext, messagebox
import threading
import os
import time
from pathlib import Path
import customtkinter as ctk  # Using customtkinter to create a more attractive UI
from CTkMessagebox import CTkMessagebox  # More attractive message box
from PIL import Image, ImageTk  # For handling icons

# Import existing summarization functions
import pdfplumber
from docx import Document
from typing import List
import subprocess


# —— 1. Text Extraction Functions —— #
def extract_text(file_path: str) -> str:
    """Automatically call the appropriate extraction function based on file extension"""
    if file_path.lower().endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    if file_path.lower().endswith(".docx"):
        return extract_text_from_docx(file_path)
    # Default to plain text
    with open(file_path, encoding="utf-8") as f:
        return f.read()


def extract_text_from_pdf(pdf_path: str) -> str:
    texts = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            texts.append(page.extract_text() or "")
    return "\n".join(texts)


def extract_text_from_docx(docx_path: str) -> str:
    doc = Document(docx_path)
    return "\n".join([p.text for p in doc.paragraphs if p.text.strip()])


# —— 2. Text Chunking Functions —— #
def chunk_text(text: str, max_chars: int = 20000) -> List[str]:
    """Split long text into characters to avoid prompt words that are too long"""
    paras = text.split("\n")
    chunks: List[str] = []
    current = []
    count = 0
    for p in paras:
        l = len(p)
        if count + l > max_chars and current:
            chunks.append("\n".join(current))
            current = [p]
            count = l
        else:
            current.append(p)
            count += l
    if current:
        chunks.append("\n".join(current))
    return chunks


# —— 3. Summarize with Ollama CLI —— #
def summarize_with_ollama(text: str, model: str) -> str:
    prompt = (
        "The following is the summary content：\n"
        f"{text}\n\n"
        "Please extract the key points of this passage in English and output a concise summary："
    )
    # Call Ollama CLI and specify decoding parameters
    proc = subprocess.run(
        ["ollama", "run", model, prompt],
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
        encoding="utf-8",
        errors="ignore",
        check=False
    )

    # Prioritize stdout; if empty, use stderr
    output = proc.stdout.strip() if proc.stdout else proc.stderr.strip()
    return output


def summarize(text: str, model: str, update_callback=None) -> str:
    chunks = chunk_text(text)
    summaries = []
    for idx, chunk in enumerate(chunks, 1):
        if update_callback:
            update_callback(f"Summarizing part {idx}/{len(chunks)}...")
        s = summarize_with_ollama(chunk, model)
        summaries.append(s)
    return "\n\n".join(summaries)


# GUI Application Class
class SummarizationApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        self.title("Intelligent Text Summarization Tool")
        self.geometry("1000x700")
        self.minsize(800, 600)

        # Set theme
        ctk.set_appearance_mode("system")  # System mode, automatically adapts to dark/light
        ctk.set_default_color_theme("blue")  # Default color theme

        self.file_path = None
        self.summary_result = None

        self.create_widgets()
        self.center_window()

    def center_window(self):
        """Center the window on display"""
        self.update_idletasks()
        width = self.winfo_width()
        height = self.winfo_height()
        x = (self.winfo_screenwidth() // 2) - (width // 2)
        y = (self.winfo_screenheight() // 2) - (height // 2)
        self.geometry(f'{width}x{height}+{x}+{y}')

    def create_widgets(self):
        # Create grid layout
        self.grid_columnconfigure(0, weight=1)
        self.grid_rowconfigure(3, weight=1)

        # Title
        header_frame = ctk.CTkFrame(self, corner_radius=0, fg_color="transparent")
        header_frame.grid(row=0, column=0, sticky="ew", padx=20, pady=(20, 10))

        title_label = ctk.CTkLabel(
            header_frame,
            text="Intelligent Text Summarization Generator",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=10)

        subtitle_label = ctk.CTkLabel(
            header_frame,
            text="Upload PDF or DOC files to quickly generate high-quality text summaries",
            font=ctk.CTkFont(size=14),
            text_color="gray"
        )
        subtitle_label.pack()

        # File selection area
        file_frame = ctk.CTkFrame(self)
        file_frame.grid(row=1, column=0, padx=20, pady=10, sticky="ew")

        self.file_label = ctk.CTkLabel(
            file_frame,
            text="No file selected",
            font=ctk.CTkFont(size=14),
            width=400
        )
        self.file_label.pack(side=tk.LEFT, padx=20, pady=15, fill=tk.X, expand=True)

        self.browse_button = ctk.CTkButton(
            file_frame,
            text="Browse Files",
            command=self.browse_file,
            width=120,
            height=35,
            font=ctk.CTkFont(size=14)
        )
        self.browse_button.pack(side=tk.RIGHT, padx=20, pady=15)

        # Control area
        control_frame = ctk.CTkFrame(self)
        control_frame.grid(row=2, column=0, padx=20, pady=10, sticky="ew")

        # Model selection dropdown
        model_label = ctk.CTkLabel(
            control_frame,
            text="Select Model:",
            font=ctk.CTkFont(size=14)
        )
        model_label.pack(side=tk.LEFT, padx=(20, 5), pady=15)

        self.model_var = tk.StringVar(value="deepseek-r1:14b")
        models = ["deepseek-r1:14b", "llama2:7b", "mistral:7b", "mixtral:8x7b"]
        self.model_combo = ctk.CTkComboBox(
            control_frame,
            values=models,
            variable=self.model_var,
            width=180,
            height=35,
            font=ctk.CTkFont(size=14)
        )
        self.model_combo.pack(side=tk.LEFT, padx=5, pady=15)

        # Summarize button
        self.summarize_button = ctk.CTkButton(
            control_frame,
            text="Start Summarizing",
            command=self.start_summarization,
            width=140,
            height=35,
            font=ctk.CTkFont(size=14, weight="bold")
        )
        self.summarize_button.pack(side=tk.RIGHT, padx=20, pady=15)

        # Export button
        self.export_button = ctk.CTkButton(
            control_frame,
            text="Export Summary",
            command=self.export_summary,
            width=140,
            height=35,
            state="disabled",
            font=ctk.CTkFont(size=14)
        )
        self.export_button.pack(side=tk.RIGHT, padx=5, pady=15)

        # Status label
        self.status_label = ctk.CTkLabel(
            control_frame,
            text="Ready",
            font=ctk.CTkFont(size=14),
            text_color="gray"
        )
        self.status_label.pack(side=tk.RIGHT, padx=20, pady=15)

        # Create tab control
        self.tab_view = ctk.CTkTabview(self)
        self.tab_view.grid(row=3, column=0, padx=20, pady=(10, 20), sticky="nsew")

        # Add summary and original text tabs
        self.tab_view.add("Summary Results")
        self.tab_view.add("Original Text")

        # Summary text box
        self.summary_text = ctk.CTkTextbox(
            self.tab_view.tab("Summary Results"),
            font=ctk.CTkFont(family="Arial", size=14),
            wrap="word"
        )
        self.summary_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Original text box
        self.original_text = ctk.CTkTextbox(
            self.tab_view.tab("Original Text"),
            font=ctk.CTkFont(family="Arial", size=14),
            wrap="word"
        )
        self.original_text.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Status bar
        status_bar = ctk.CTkFrame(self, height=25, fg_color=("gray90", "gray20"))
        status_bar.grid(row=4, column=0, sticky="ew")

        status_text = ctk.CTkLabel(
            status_bar,
            text="© 2024 Intelligent Text Summarization Tool",
            font=ctk.CTkFont(size=12),
            text_color="gray"
        )
        status_text.pack(side=tk.RIGHT, padx=10)

        # Progress indicator
        self.progress_bar = ctk.CTkProgressBar(self)
        self.progress_bar.grid(row=5, column=0, padx=20, pady=(0, 20), sticky="ew")
        self.progress_bar.set(0)

    def browse_file(self):
        """Open file dialog to select file"""
        file_types = [
            ("Document Files", "*.pdf;*.docx;*.txt"),
            ("PDF Files", "*.pdf"),
            ("Word Files", "*.docx"),
            ("Text Files", "*.txt"),
            ("All Files", "*.*")
        ]
        file_path = filedialog.askopenfilename(filetypes=file_types)

        if file_path:
            self.file_path = file_path
            filename = os.path.basename(file_path)
            self.file_label.configure(text=f"Selected: {filename}")

            # Clear text boxes
            self.summary_text.delete("0.0", tk.END)
            self.original_text.delete("0.0", tk.END)

            # Disable export button
            self.export_button.configure(state="disabled")

            # Update status
            self.status_label.configure(text="File loaded")

            # Load original text
            try:
                raw_text = extract_text(file_path)
                self.original_text.insert("0.0", raw_text)
                # Show a small preview of the original text
                preview = raw_text[:500] + "..." if len(raw_text) > 500 else raw_text
                self.summary_text.insert("0.0", f"Original text preview:\n{preview}\n\nClick 'Start Summarizing' to generate summary.")

                # Activate progress bar animation
                self.animate_progress(0, 0.3, 10)
            except Exception as e:
                self.show_error(f"Error loading file: {str(e)}")

    def start_summarization(self):
        """Start the summarization process"""
        if not self.file_path:
            self.show_error("Please select a file first")
            return

        # Disable buttons to prevent repeated clicks
        self.summarize_button.configure(state="disabled")
        self.browse_button.configure(state="disabled")

        # Clear summary text box
        self.summary_text.delete("0.0", tk.END)
        self.summary_text.insert("0.0", "Generating summary, please wait...\n")

        # Reset and start progress bar
        self.progress_bar.set(0.3)
        self.animate_progress(0.3, 0.9, duration=0.05, repeat=True)

        # Update status
        self.status_label.configure(text="Processing...")

        # Run summarization process in a separate thread
        threading.Thread(target=self.run_summarization, daemon=True).start()

    def animate_progress(self, start_val, end_val, duration, step=0.01, repeat=False):
        """Animate progress bar display"""

        def update_progress(current):
            if not self.winfo_exists():
                return

            if current <= end_val:
                self.progress_bar.set(current)
                increment = step
                self.after(int(duration * 1000),
                           lambda: update_progress(current + increment))
            elif repeat:
                self.progress_bar.set(start_val)
                self.after(int(duration * 1000),
                           lambda: update_progress(start_val + step))

        update_progress(start_val)

    def run_summarization(self):
        """Run summarization process in background thread"""
        try:
            # Get model name
            model = self.model_var.get()

            # Get original text
            raw_text = extract_text(self.file_path)

            # Define update callback function
            def update_status(msg):
                self.update_ui(lambda: self.summary_text.insert(tk.END, f"\n{msg}"))

            # Run summarization
            summary = summarize(raw_text, model, update_callback=update_status)
            self.summary_result = summary

            # Update UI
            self.update_ui(lambda: self.update_summary_result(summary))

        except Exception as e:
            self.update_ui(lambda: self.show_error(f"Error generating summary: {str(e)}"))
        finally:
            # Restore UI state
            self.update_ui(self.restore_ui_state)

    def update_ui(self, func):
        """Update UI in main thread"""
        if self.winfo_exists():
            self.after(0, func)

    def update_summary_result(self, summary):
        """Update summary result"""
        # Clear and display result
        self.summary_text.delete("0.0", tk.END)
        self.summary_text.insert("0.0", summary)

        # Enable export button
        self.export_button.configure(state="normal")

        # Update status
        self.status_label.configure(text="Summary completed")

        # Switch to summary tab
        self.tab_view.set("Summary Results")

        # Complete progress bar
        self.progress_bar.set(1.0)
        self.after(1000, lambda: self.progress_bar.set(0))

    def restore_ui_state(self):
        """Restore UI state"""
        self.summarize_button.configure(state="normal")
        self.browse_button.configure(state="normal")

    def export_summary(self):
        """Export summary to file"""
        if not self.summary_result:
            self.show_error("No summary available to export")
            return

        # Get original filename (without extension)
        if self.file_path:
            original_filename = Path(self.file_path).stem
            default_filename = f"{original_filename}_summary.txt"
        else:
            default_filename = "summary.txt"

        # Show save dialog
        file_types = [
            ("Text Files", "*.txt"),
            ("Markdown Files", "*.md"),
            ("Word Documents", "*.docx"),
            ("All Files", "*.*")
        ]
        save_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=file_types,
            initialfile=default_filename
        )

        if save_path:
            try:
                # Decide how to save based on file extension
                ext = Path(save_path).suffix.lower()

                if ext == '.docx':
                    # Save as Word document
                    doc = Document()
                    doc.add_heading('Document Summary', 0)

                    # Add original file information
                    if self.file_path:
                        doc.add_paragraph(f'Original file: {os.path.basename(self.file_path)}')

                    doc.add_paragraph('')  # Empty line

                    # Add summary content
                    for para in self.summary_result.split('\n'):
                        if para.strip():
                            doc.add_paragraph(para)

                    doc.save(save_path)
                else:
                    # Save as text file
                    with open(save_path, 'w', encoding='utf-8') as f:
                        # Add title
                        f.write("# Document Summary\n\n")

                        # Add original file information
                        if self.file_path:
                            f.write(f"Original file: {os.path.basename(self.file_path)}\n\n")

                        # Add summary content
                        f.write(self.summary_result)

                # Show success message
                CTkMessagebox(
                    title="Export Successful",
                    message=f"Summary successfully exported to:\n{save_path}",
                    icon="check",
                    option_1="OK"
                )

                # Update status
                self.status_label.configure(text="Summary exported")

            except Exception as e:
                self.show_error(f"Error exporting summary: {str(e)}")

    def show_error(self, message):
        """Show error message"""
        CTkMessagebox(
            title="Error",
            message=message,
            icon="cancel",
            option_1="OK"
        )
        self.status_label.configure(text="Error occurred")


# Main program entry
if __name__ == "__main__":
    app = SummarizationApp()
    app.mainloop()